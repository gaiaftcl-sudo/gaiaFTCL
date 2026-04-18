#!/usr/bin/env python3
from __future__ import annotations

import argparse
import glob
import hashlib
import json
import pathlib
import subprocess
import sys
import zipfile
import re


def pick_latest(pattern: str) -> str | None:
    files = sorted(glob.glob(pattern))
    return files[-1] if files else None


def read_json_if_exists(p: pathlib.Path) -> dict | None:
    if not p.exists():
        return None
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return None


_XML_FORBIDDEN_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F]")


def xml_safe_text(value: object) -> str:
    s = str(value)
    return _XML_FORBIDDEN_RE.sub("", s)


def add_p(doc, text: object) -> None:
    doc.add_paragraph(xml_safe_text(text))


def append_c4_semantics_docx(doc, text: str) -> None:
    """Insert RELEASE_C4_SEMANTICS.md: headings + paragraphs (no external MD parser)."""
    for raw in text.split("\n\n"):
        block = raw.strip()
        if not block:
            continue
        lines = block.split("\n")
        head = lines[0].strip()
        tail = "\n".join(lines[1:]).strip()
        if head.startswith("### "):
            doc.add_heading(head[4:].strip(), 3)
            if tail:
                for ln in tail.split("\n"):
                    if ln.strip():
                        add_p(doc, ln.strip())
        elif head.startswith("## "):
            doc.add_heading(head[3:].strip(), 2)
            if tail:
                for ln in tail.split("\n"):
                    if ln.strip():
                        add_p(doc, ln.strip())
        elif head.startswith("# "):
            doc.add_heading(head[2:].strip(), 1)
            if tail:
                for ln in tail.split("\n"):
                    if ln.strip():
                        add_p(doc, ln.strip())
        else:
            add_p(doc, block[:15000])


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--repo-root", type=pathlib.Path, default=pathlib.Path(__file__).resolve().parents[1])
    ap.add_argument("--report-json", "--json", dest="report_json", type=pathlib.Path, required=True)
    ap.add_argument("--report-md", type=pathlib.Path)
    ap.add_argument("--out-docx", "--output", dest="out_docx", type=pathlib.Path, required=True)
    ap.add_argument(
        "--c4-semantics",
        dest="c4_semantics",
        type=pathlib.Path,
        default=None,
        help="Markdown: Birth vs Earth semantics (default: evidence/discord/RELEASE_C4_SEMANTICS.md)",
    )
    ap.add_argument(
        "--require-c4-semantics",
        action="store_true",
        help="REFUSED if semantics file missing or C4 headings not present in written DOCX XML",
    )
    args = ap.parse_args()

    if args.require_c4_semantics:
        mat = subprocess.run(
            [
                sys.executable,
                str(args.repo_root / "scripts" / "materialize_c4_semantics_md.py"),
                "--repo-root",
                str(args.repo_root),
            ],
            capture_output=True,
            text=True,
        )
        if mat.returncode != 0:
            raise SystemExit(
                "REFUSED: materialize_c4_semantics_md failed — "
                + (mat.stderr or mat.stdout or "").strip()[:800]
            )

    try:
        from docx import Document
        from docx.shared import Inches
    except Exception as exc:
        raise SystemExit(f"REFUSED: python-docx not available: {exc}")

    if not args.report_md:
        md_guess = args.report_json.with_suffix(".md")
        if md_guess.exists():
            args.report_md = md_guess
        else:
            fallback = pick_latest(str(args.repo_root / "evidence" / "discord" / "RELEASE_REPORT_*.md"))
            if fallback:
                args.report_md = pathlib.Path(fallback)
            else:
                args.report_md = None

    report = {}
    if args.report_json.exists():
        report = json.loads(args.report_json.read_text(encoding="utf-8"))
    if args.report_md and args.report_md.exists():
        md_text = args.report_md.read_text(encoding="utf-8")
    else:
        md_text = "# Release Report (fallback)\n\n" + json.dumps(report, indent=2, ensure_ascii=True)[:12000]

    doc = Document()
    doc.add_heading("GaiaFTCL Release Verification Bundle", 0)
    verdict = (
        "VERIFIED"
        if report.get("state") == "CALORIE"
        and report.get("uniformity") == "UNIFORM"
        and int(report.get("failed_steps", 0)) == 0
        else "NOT VERIFIED"
    )
    add_p(doc, f"C4 Matrix Verdict: {verdict}")
    add_p(
        doc,
        f"State: {report.get('state')} | Uniformity: {report.get('uniformity')} | Failed steps: {report.get('failed_steps', 0)}"
    )

    sem_path = args.c4_semantics or (args.repo_root / "evidence" / "discord" / "RELEASE_C4_SEMANTICS.md")
    if args.require_c4_semantics and not sem_path.exists():
        raise SystemExit(f"REFUSED: --require-c4-semantics but file missing: {sem_path}")
    if sem_path.exists():
        doc.add_heading("C4 Full Release Semantics (Social vs Planetary)", level=1)
        sem_bytes = sem_path.read_bytes()
        sem_md5 = hashlib.md5(sem_bytes).hexdigest()
        append_c4_semantics_docx(doc, sem_bytes.decode(encoding="utf-8", errors="replace"))
        add_p(doc, f"Source: {sem_path}")
        add_p(doc, f"C4_SEMANTICS_MD5: {sem_md5}")
    else:
        add_p(doc, f"WARN: C4 semantics file missing (expected {sem_path})")

    doc.add_heading("Release Report (Markdown Witness)", level=1)
    add_p(doc, md_text[:6000])

    doc.add_heading("Language Games", level=1)
    lg = pick_latest(str(args.repo_root / "evidence" / "discord" / "LANGUAGE_GAMES_*.json"))
    if lg:
        games = json.loads(pathlib.Path(lg).read_text(encoding="utf-8")).get("games", [])
        for g in games:
            add_p(
                doc,
                f"{g.get('domain')} {g.get('command')} kind={g.get('kind')} source={g.get('source')} receipt={g.get('receipt_id')}"
            )

    validation_json = args.repo_root / "evidence" / "discord" / "LATEST_SCREENSHOT_VALIDATION.json"
    rc = subprocess.run(
        [
            "python3",
            str(args.repo_root / "scripts" / "validate_screenshots.py"),
            "--repo-root",
            str(args.repo_root),
            "--out-json",
            str(validation_json),
        ],
        check=False,
        capture_output=True,
        text=True,
    )
    if rc.returncode != 0:
        # Self-heal fallback: still emit the report document, but mark screenshot evidence missing.
        print(f"WARN: screenshot validation failed, continuing without embeds: {rc.stderr or rc.stdout}")
        validation = {"total": 0, "passed": 0, "failed": 0, "valid_screenshots": []}
    else:
        validation = json.loads(validation_json.read_text(encoding="utf-8"))
    doc.add_heading("Screenshots (Validated Before Embed)", level=1)
    add_p(
        doc,
        f"Validated screenshots: {validation.get('passed')} / {validation.get('total')} (failed: {validation.get('failed')})"
    )
    if not validation.get("valid_screenshots"):
        add_p(doc, "No validated screenshots were available in this run.")
    for shot in validation.get("valid_screenshots", []):
        p = pathlib.Path(shot)
        if p.exists() and p.stat().st_size > 0:
            add_p(doc, str(p))
            try:
                doc.add_picture(str(p), width=Inches(6.5))
            except Exception:
                add_p(doc, "(image embedding skipped for unsupported format)")

    doc.add_heading("Dual-User Convergence", level=1)
    dual_latest = pick_latest(str(args.repo_root / "evidence" / "discord" / "dual_user" / "*" / "DUAL_USER_WITNESS.json"))
    if not dual_latest:
        add_p(doc, "No dual-user witness found for this run.")
    else:
        dpath = pathlib.Path(dual_latest)
        d = read_json_if_exists(dpath) or {}
        add_p(doc, f"Witness: {dpath}")
        ua = d.get("user_a", {})
        ub = d.get("user_b", {})
        add_p(doc, f"User A source={ua.get('source')} release_id={ua.get('release_id')}")
        add_p(doc, f"User B source={ub.get('source')} release_id={ub.get('release_id')}")
        add_p(doc, f"Convergence ms={d.get('convergence_ms')}")
        if d.get("protocol_note"):
            add_p(doc, str(d.get("protocol_note")))
        mp = d.get("mooring_phases")
        if isinstance(mp, dict):
            add_p(doc, f"Onboarding / moor phases: {json.dumps(mp, ensure_ascii=True)[:3500]}")
        c = d.get("criteria", {})
        add_p(
            doc,
            "Criteria: "
            f"release_id_match={c.get('state_convergence_release_id')} | "
            f"source_diversity={c.get('source_diversity')} | "
            f"lt_2s={c.get('convergence_lt_2s')} | "
            f"earth_full_moor={c.get('earth_full_moor')} | "
            f"wallet_moored={c.get('wallet_moored')}"
        )
        for shot_name in ("user-a-phase1-onboarding.png", "user-a-moorer.png", "user-b-observer.png"):
            sp = dpath.parent / shot_name
            if sp.exists() and sp.stat().st_size > 0:
                add_p(doc, str(sp))
                try:
                    doc.add_picture(str(sp), width=Inches(6.5))
                except Exception:
                    add_p(doc, "(image embedding skipped for unsupported format)")

    args.out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.out_docx))
    if args.require_c4_semantics:
        with zipfile.ZipFile(args.out_docx, "r") as zf:
            xml = zf.read("word/document.xml").decode("utf-8", errors="replace")
        if (
            "C4 Full Release Semantics" not in xml
            or "Social vs Planetary" not in xml
            or "C4_SEMANTICS_MD5:" not in xml
        ):
            raise SystemExit(
                "REFUSED: --require-c4-semantics DOCX verification failed (missing C4 headings or C4_SEMANTICS_MD5 in word/document.xml)"
            )
    # Not a governor terminal — avoid confusion with gate-local CALORIE / process exit.
    print(f"SEALING: wrote docx {args.out_docx}", flush=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
